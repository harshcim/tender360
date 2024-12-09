import PyPDF2 # type: ignore
import os
import sys
from docx import Document

sys.path.append(os.path.join(os.path.dirname(os.path.abspath(__file__)),".."))

from log.logger import setup_logger

logger = setup_logger("pdf_extractor_logs")

def extract_text_from_pdf(pdf_path):
    
    logger.info(f"Extracting text from PDF: {pdf_path}")
    
    
    tender_documents = []
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                text = page.extract_text()
                if text:
                    tender_documents.append({
                        'page': page_num + 1,  # Page numbers start at 1
                        'text': text
                    })
                
            logger.debug(f"Extracted text from {num_pages} pages.")
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        
    return tender_documents

def extract_text_from_docx(docx_path):
    logger.info(f"Extracting text from DOCX: {docx_path}")
    
    tender_documents = []
    try:
        document = Document(docx_path)
        for idx, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if text:
                tender_documents.append({
                    'page': idx + 1,  # Using paragraph index as a "page" indicator
                    'text': text
                })
                
        logger.debug(f"Extracted text from {len(document.paragraphs)} paragraphs.")
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
    
    return tender_documents

def extract_text_from_file(file_path):
    if file_path.lower().endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif file_path.lower().endswith(".docx"):
        return extract_text_from_docx(file_path)
    else:
        logger.error(f"Unsupported file format: {file_path}")
        return []